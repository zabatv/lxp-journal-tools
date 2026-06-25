import io
import os
import json
import logging
from datetime import datetime, timezone
from concurrent.futures import ThreadPoolExecutor, as_completed

import requests
from docx import Document
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import HTMLResponse, Response
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel
from starlette.requests import Request

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

app = FastAPI(title="LXP Journal Filler")
templates = Jinja2Templates(directory="templates")

API_URL = "https://api.newlxp.ru/graphql"
GRADE_MAP = {"TWO": "2", "THREE": "3", "FOUR": "4", "FIVE": "5"}


# ---------- GraphQL Queries ----------

QUERY_SIGN_IN = """
    query SignIn($input: SignInInput!) {
        signIn(input: $input) { accessToken }
    }
"""

QUERY_GET_ME = """
    query {
        getMe {
            assignedSuborganizations {
                suborganizationId
                suborganization { id name organizationId }
            }
        }
    }
"""

QUERY_STUDY_PERIODS = """
    query {{
        studyPeriods(input: {{ filters: {{ organizationId: "{org_id}" }} }}) {{
            id name startDate endDate
        }}
    }}
"""

QUERY_GROUPS_BY_PERIOD = """
    query {{
        learningGroupsByStudyPeriodIdAndSuborganizationId(input: {{
            studyPeriodId: "{study_period_id}"
            suborganizationId: "{suborg_id}"
        }}) {{ id name }}
    }}
"""

QUERY_GROUPS_BY_ORG = """
    query {{
        getLearningGroups(input: {{ 
            organizationId: "{org_id}" 
            suborganizationId: "{suborg_id}" 
            isArchived: false 
        }}) {{
            id name
        }}
    }}
"""

QUERY_DISCIPLINES = """
    query {{
        disciplinesByGroups(input: {{ groupIds: ["{group_id}"] }}) {{
            id name code
            teachers {{ user {{ lastName firstName middleName }} }}
        }}
    }}
"""

QUERY_STUDENTS = """
    query {{
        searchStudentsInLearningGroup(input: {{
            filters: {{ learningGroupId: "{group_id}", isExpelled: false }}
        }}) {{
            items {{ id user {{ lastName firstName middleName }} }}
        }}
    }}
"""

QUERY_STUDENT_DISCIPLINES = """
    query {{
        searchStudentDisciplines(input: {{
            studentId: "{student_id}"
            filters: {{ studyPeriodId: "{study_period_id}" }}
        }}) {{
            disciplineId
            disciplineGrade
            disciplineGrade_V2
            scoreForAnsweredTasks
            maxScoreForAnsweredTasks
            hasRetake
            retakeDisciplineGrade
            retakeScore
            topics {{
                ... on StudentTopic {{
                    status
                    topicScore
                    topic {{
                        id
                        name
                    }}
                }}
            }}
        }}
    }}
"""

QUERY_USER_GRADE = """
    query {{
        getUserById(input: {{ userId: "{student_id}" }}) {{
            student {{ 
                studentDiscipline(disciplineId: "{disc_id}") {{ 
                    disciplineGrade
                    disciplineGrade_V2
                    scoreForAnsweredTasks
                    maxScoreForAnsweredTasks
                    hasRetake
                    retakeDisciplineGrade
                    retakeScore
                    topics {{
                        ... on StudentTopic {{
                            status
                            topicScore
                            topic {{
                                id
                                name
                            }}
                        }}
                    }}
                }}
            }}
        }}
    }}
"""


# ---------- GraphQL helper ----------

def graphql(token: str, query: str, variables: dict = None, timeout: int = 30) -> dict:
    clean_token = token[7:] if token.startswith("Bearer ") else token
    headers = {"Authorization": f"Bearer {clean_token}", "Content-Type": "application/json"}
    body = {"query": query}
    if variables:
        body["variables"] = variables
    
    logger.info(f"=== GraphQL Request ===")
    logger.info(f"Query: {query.strip()}")
    if variables:
        logger.info(f"Variables: {variables}")
    
    try:
        resp = requests.post(API_URL, headers=headers, json=body, timeout=timeout)
        logger.info(f"Response status: {resp.status_code}")
        
        if resp.status_code != 200:
            logger.error(f"Response body: {resp.text}")
            resp.raise_for_status()
        
        data = resp.json()
        
        if data.get("errors"):
            logger.error(f"GraphQL errors: {data['errors']}")
            raise HTTPException(status_code=400, detail=data["errors"][0].get("message", "GraphQL error"))
        
        return data.get("data", {})
        
    except requests.RequestException as e:
        logger.error(f"Request failed: {e}")
        raise HTTPException(status_code=502, detail=f"LXP недоступен: {e}")


# ---------- Pages ----------

@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    return templates.TemplateResponse("login.html", {"request": request})


@app.get("/dashboard", response_class=HTMLResponse)
async def dashboard(request: Request):
    return templates.TemplateResponse("dashboard.html", {"request": request})


# ---------- Auth ----------

class LoginInput(BaseModel):
    email: str
    password: str


@app.post("/api/auth/login")
async def login(input: LoginInput):
    data = graphql("", QUERY_SIGN_IN, variables={"input": {"email": input.email.strip(), "password": input.password}})
    token = data.get("signIn", {}).get("accessToken")
    if not token:
        raise HTTPException(status_code=401, detail="Токен не получен")
    return {"token": token}


@app.post("/api/auth/check")
async def check_token(request: Request):
    body = await request.json()
    token = body.get("token", "")
    if not token:
        raise HTTPException(status_code=400, detail="Токен не предоставлен")
    try:
        graphql(token, "query { getMe { id } }")
        return {"valid": True}
    except HTTPException:
        raise HTTPException(status_code=401, detail="Токен недействителен")


# ---------- Data API ----------

@app.get("/api/suborganizations")
async def get_suborganizations(token: str = ""):
    if not token:
        raise HTTPException(status_code=401, detail="Требуется токен")
    data = graphql(token, QUERY_GET_ME)
    items = data["getMe"]["assignedSuborganizations"]
    seen = set()
    result = []
    for item in items:
        key = item["suborganizationId"]
        if key not in seen:
            seen.add(key)
            result.append(item)
    return {"items": result}


@app.get("/api/study-periods")
async def get_study_periods(token: str = "", org_id: str = ""):
    if not token:
        raise HTTPException(status_code=401, detail="Требуется токен")
    
    query = QUERY_STUDY_PERIODS.format(org_id=org_id)
    data = graphql(token, query)
    
    now = datetime.now(timezone.utc)
    items = sorted(data["studyPeriods"], key=lambda x: x["startDate"])
    
    for sp in items:
        try:
            end = datetime.fromisoformat(sp["endDate"].replace("Z", "+00:00"))
            start = datetime.fromisoformat(sp["startDate"].replace("Z", "+00:00"))
            sp["isCurrent"] = start <= now <= end
        except Exception:
            sp["isCurrent"] = False
    
    return {"items": items}


@app.get("/api/groups")
async def get_groups(token: str = "", org_id: str = "", suborg_id: str = "", study_period_id: str = ""):
    if not token:
        raise HTTPException(status_code=401, detail="Требуется токен")
    
    if study_period_id:
        query = QUERY_GROUPS_BY_PERIOD.format(study_period_id=study_period_id, suborg_id=suborg_id)
        data = graphql(token, query)
        return {"items": data["learningGroupsByStudyPeriodIdAndSuborganizationId"]}
    
    query = QUERY_GROUPS_BY_ORG.format(org_id=org_id, suborg_id=suborg_id)
    data = graphql(token, query)
    return {"items": data["getLearningGroups"]}


@app.get("/api/disciplines")
async def get_disciplines(token: str = "", group_id: str = ""):
    if not token:
        raise HTTPException(status_code=401, detail="Требуется токен")
    query = QUERY_DISCIPLINES.format(group_id=group_id)
    data = graphql(token, query)
    return {"items": data["disciplinesByGroups"]}


def _determine_grade_from_topics(sd: dict, student_id: str = "") -> dict:
    """
    Определяет оценку и наличие пересдачи.
    Логика:
    - Берём disciplineGrade_V2 как основную оценку
    - hasRetake берём из API, а не вычисляем из FAILED тем
    - retakeDisciplineGrade используем только если hasRetake: True
    """
    result = {
        "grade": "",
        "hasRetake": False,
        "retakeGrade": "",
        "retakeScore": "",
    }
    
    # Логируем сырые данные для отладки
    logger.info(f"=== Student {student_id} ===")
    logger.info(f"disciplineGrade: {sd.get('disciplineGrade')}")
    logger.info(f"disciplineGrade_V2: {sd.get('disciplineGrade_V2')}")
    logger.info(f"hasRetake: {sd.get('hasRetake')}")
    logger.info(f"retakeDisciplineGrade: {sd.get('retakeDisciplineGrade')}")
    
    # Берём disciplineGrade_V2 как основную оценку
    grade_v2 = sd.get("disciplineGrade_V2") or sd.get("disciplineGrade") or ""
    if grade_v2 in GRADE_MAP:
        result["grade"] = GRADE_MAP[grade_v2]
    elif grade_v2:
        result["grade"] = grade_v2
    
    # hasRetake берём из API
    has_retake = sd.get("hasRetake", False)
    result["hasRetake"] = has_retake
    
    # retakeDisciplineGrade используем только если hasRetake: True
    if has_retake:
        retake_grade = sd.get("retakeDisciplineGrade")
        if retake_grade and retake_grade in GRADE_MAP:
            result["retakeGrade"] = GRADE_MAP[retake_grade]
            # Если есть валидная оценка за пересдачу, она приоритетнее
            result["grade"] = GRADE_MAP[retake_grade]
        elif retake_grade:
            result["retakeGrade"] = retake_grade
    
    logger.info(f"Final grade: {result['grade']}, hasRetake: {result['hasRetake']}, retakeGrade: {result['retakeGrade']}")
    logger.info("=" * 50)
    
    return result


@app.get("/api/students")
async def get_students(token: str = "", group_id: str = "", disc_id: str = "", study_period_id: str = ""):
    if not token:
        raise HTTPException(status_code=401, detail="Требуется токен")

    query1 = QUERY_STUDENTS.format(group_id=group_id)
    students_data = graphql(token, query1)
    students = students_data["searchStudentsInLearningGroup"]["items"]

    query2 = QUERY_DISCIPLINES.format(group_id=group_id)
    disc_data = graphql(token, query2)

    teacher_name = ""
    for d in disc_data["disciplinesByGroups"]:
        if d["id"] == disc_id and d["teachers"]:
            t = d["teachers"][0]["user"]
            teacher_name = f"{t['lastName']} {t['firstName']} {t.get('middleName', '')}".strip()
            break

    name_map = {
        s["id"]: f"{s['user']['lastName']} {s['user']['firstName']} {s['user'].get('middleName', '')}".strip()
        for s in students
    }

    def get_grade(student_id: str, idx: int) -> dict:
        base = {
            "id": student_id,
            "name": name_map.get(student_id, "Ошибка"),
            "grade": "",
            "hasRetake": False,
            "retakeGrade": "",
            "retakeScore": "",
            "idx": idx,
        }
        try:
            if study_period_id:
                query3 = QUERY_STUDENT_DISCIPLINES.format(student_id=student_id, study_period_id=study_period_id)
                sd_data = graphql(token, query3)
                for sd in sd_data["searchStudentDisciplines"]:
                    if sd["disciplineId"] == disc_id:
                        grade_info = _determine_grade_from_topics(sd, student_id)
                        base.update(grade_info)
                        break
            else:
                query4 = QUERY_USER_GRADE.format(student_id=student_id, disc_id=disc_id)
                gdata = graphql(token, query4)
                sd = gdata["getUserById"]["student"]["studentDiscipline"]
                if sd:
                    grade_info = _determine_grade_from_topics(sd, student_id)
                    base.update(grade_info)
        except Exception as e:
            logger.error(f"Error fetching grade for student {student_id}: {e}")
        return base

    results = []
    with ThreadPoolExecutor(max_workers=10) as executor:
        futures = {executor.submit(get_grade, s["id"], i): s for i, s in enumerate(students)}
        for future in as_completed(futures):
            try:
                results.append(future.result(timeout=15))
            except Exception as e:
                logger.error(f"Timeout or error in thread: {e}")
                s = futures[future]
                results.append({
                    "id": s["id"],
                    "name": name_map.get(s["id"], "Ошибка"),
                    "grade": "",
                    "hasRetake": False,
                    "retakeGrade": "",
                    "retakeScore": "",
                    "idx": s["id"],
                })
    results.sort(key=lambda x: x["idx"])

    return {"items": results, "teacher_name": teacher_name, "count": len(results)}


# ---------- DOCX Fill ----------

def _replace_in_paragraph(paragraph, mapping: dict) -> None:
    full_text = "".join(run.text for run in paragraph.runs)
    if not any(ph in full_text for ph in mapping):
        return

    new_text = full_text
    for placeholder, value in mapping.items():
        new_text = new_text.replace(placeholder, value)

    if new_text == full_text:
        return

    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


def _replace_in_cell(cell, mapping: dict) -> None:
    for para in cell.paragraphs:
        _replace_in_paragraph(para, mapping)


@app.post("/api/docx/fill")
async def fill_docx(
    file: UploadFile = File(...),
    token: str = Form(...),
    group_name: str = Form(""),
    disc_name: str = Form(""),
    teacher_name: str = Form(""),
    students_json: str = Form("[]"),
):
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Файл должен быть в формате DOCX")

    content = await file.read()
    doc = Document(io.BytesIO(content))
    students = json.loads(students_json)

    global_mapping = {"%g": group_name, "%d": disc_name, "%t": teacher_name}
    for para in doc.paragraphs:
        _replace_in_paragraph(para, global_mapping)

    student_idx = 0
    for table in doc.tables:
        for row in table.rows:
            row_text = "".join(cell.text for cell in row.cells)
            has_student_marker = "%n" in row_text or "%q" in row_text

            for cell in row.cells:
                if has_student_marker and student_idx < len(students):
                    student = students[student_idx]
                    grade = str(student.get("grade") or "") or "—"
                    mapping = {
                        "%g": group_name,
                        "%d": disc_name,
                        "%t": teacher_name,
                        "%n": student.get("name", ""),
                        "%q": grade,
                    }
                else:
                    mapping = global_mapping
                _replace_in_cell(cell, mapping)

            if has_student_marker:
                student_idx += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=filled_{file.filename}"},
    )


@app.get("/example")
async def download_example():
    file_path = os.path.join(os.path.dirname(__file__), "static", "example.docx")
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="example.docx not found")
    with open(file_path, "rb") as f:
        content = f.read()
    return Response(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=example.docx"},
    )


if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    logger.info(f"Starting LXP Journal Filler on port {port}")
    uvicorn.run(app, host="0.0.0.0", port=port, log_level="info")
